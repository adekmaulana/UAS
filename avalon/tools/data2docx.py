import datetime

from typing import List, Union
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from avalon.tools import windows, rupiah_format


WINDOWS: Union[str, int] = ['CodeNewRoman NF', 10]
LINUX: Union[str, int] = ['CodeNewRoman Nerd Font Mono', 11]
document: Document


def parse_data(d: datetime.datetime, no: List[int]) -> None:
    order: str = f"#AV{no[0]}{no[1]}{no[2]}"
    cells_text: List[str, str] = [
        f"{d:%d}/{d:%m}/{d:%y} {d:%I}:{d:%M}:{d:%S} {d:%p}",
        f"{order}"
    ]
    if windows is True:
        document = Document('avalon/tools/receipt_windows.docx')
    else:
        document = Document('avalon/tools/receipt_linux.docx')
    table = document.tables[1].rows[0]

    for index, cell in enumerate(table.cells):
        paragraphs = cell.paragraphs
        for paragraph in paragraphs:
            run = paragraph.add_run()
            font = run.font
            if windows is True:
                font.name = WINDOWS[0]
                font.size = Pt(WINDOWS[1] + 1)
            else:
                font.name = LINUX[0]
                font.size = Pt(LINUX[1] + 1)
            font.bold = True
            run.text = cells_text[index]
    return document.save('avalon/tools/temp.docx')


def parse_order(list_belanja: List[str],
                total_bayar: List[int],
                data: List[int]) -> None:
    TEXT: List[str, ...] = ["SUBTOTAL", "PPN 10%", "TOTAL", "UANG", "CHANGE"]
    right = WD_ALIGN_PARAGRAPH.RIGHT

    document = Document('avalon/tools/temp.docx')
    table = document.tables[2]

    total = len(list_belanja) + 5
    for row in range(total):
        table.add_row()
    rows = table.rows

    # Tulis belanjaan pada struk belanja
    for index, belanja in enumerate(list_belanja):
        for i, cell in enumerate(rows[index].cells):
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                run = paragraph.add_run()
                font = run.font
                if windows is True:
                    font.name = WINDOWS[0]
                    font.size = Pt(WINDOWS[1])
                else:
                    font.name = LINUX[0]
                    font.size = Pt(LINUX[1])
                font.bold = True
                if i == 0:
                    paragraph.paragraph_format.alignment = right
                    run.text = list_belanja[index]
                elif i == 1:
                    run.text = rupiah_format(total_bayar[index])

    # Tulis SUBTOTAL, PPN, TOTAL, UANG, CHANGE/KEMBALIAN
    i = len(list_belanja) + 1
    j = i
    while i <= total:
        for index, cell in enumerate(rows[i].cells):
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                run = paragraph.add_run()
                font = run.font
                if windows is True:
                    font.name = WINDOWS[0]
                    font.size = Pt(WINDOWS[1])
                else:
                    font.name = LINUX[0]
                    font.size = Pt(LINUX[1])
                font.bold = True
                if index == 0:
                    paragraph.paragraph_format.alignment = right
                    run.text = TEXT[i - j]
                elif index == 1:
                    run.text = rupiah_format(data[i - j])
        i += 1

    return document.save('avalon/tools/temp.docx')


def merge(d: datetime.datetime,
          no: List[int],
          list_belanja: List[str],
          total_bayar: List[int],
          data: List[int]) -> None:
    parse_data(d, no)
    parse_order(list_belanja, total_bayar, data)
