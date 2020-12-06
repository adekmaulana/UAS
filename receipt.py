import os
import sys
import datetime
import cloudconvert
import locale

from contextlib import contextmanager
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from tools import convert

from subprocess import PIPE, Popen, CalledProcessError


WINDOWS = ['CodeNewRoman NF', 10]
LINUX = ['CodeNewRoman Nerd Font Mono', 11]

windows = False
if os.name == 'nt':
    windows = True


def subprocess_run(cmd):
    subproc = Popen(cmd, stdout=PIPE, stderr=PIPE,
                    shell=True, universal_newlines=True)
    talk = subproc.communicate()
    exitCode = subproc.returncode
    if exitCode != 0:
        raise CalledProcessError(exitCode, cmd)
    return talk


def delete_last_lines(n=1):
    CURSOR_UP_ONE = '\x1b[1A'
    ERASE_LINE = '\x1b[2K'
    for _ in range(n):
        sys.stdout.write(CURSOR_UP_ONE)
        sys.stdout.write(ERASE_LINE)
        sys.stdout.flush()


def rupiah_format(angka, with_prefix=True, desimal=0):
    locale.setlocale(locale.LC_NUMERIC, 'id_ID.UTF-8')
    rupiah = locale.format_string("%.*f", (desimal, angka), True)
    if with_prefix:
        return "Rp. {}".format(rupiah)
    return rupiah


@contextmanager
def suppress_stdout():
    """ Didapat dari https://stackoverflow.com/a/60327812 """
    with open(os.devnull, "w") as devnull:
        old_stdout = sys.stdout
        sys.stdout = devnull
        try:
            yield
        finally:
            sys.stdout = old_stdout


def cv() -> None:
    tasks = {
        "tasks": {
            "upload-file": {
                "operation": "import/upload"
            },
            "convert-file-png": {
                "operation": "convert",
                "input": "upload-file",
                "input_format": "pdf",
                "output_format": "png",
                "engine": "poppler",
                "engine_version": "0.71.0",
                "filename": "receipt.png"
            },
            "download-converted-file": {
                "operation": "export/url",
                "input": "convert-file-png"
            }
        }
    }

    # convert ke PDF
    if windows is True:
        convert('temp.docx')
    else:
        cmd = 'unoconv -f pdf temp.docx'
        talk = subprocess_run(cmd)

    # Initialized API_KEY
    cloudconvert.configure(api_key=API_KEY)

    # Buat job pada server
    job = cloudconvert.Job.create(payload=tasks)

    # Upload file ke server
    upload_task_id = job['tasks'][0]['id']
    upload_task = cloudconvert.Task.find(id=upload_task_id)
    res = cloudconvert.Task.upload(file_name='temp.pdf', task=upload_task)

    # mulai convert pada server
    convert_task_id = job['tasks'][1]['id']
    res = cloudconvert.Task.wait(id=convert_task_id)  # pdf to png

    # download converted file di server ke local
    download_task_id = job['tasks'][2]['id']
    res = cloudconvert.Task.wait(id=download_task_id)
    files = res.get("result").get("files")[0]

    if os.path.exists(files['filename']):
        os.remove(files['filename'])

    with suppress_stdout():
        res = cloudconvert.download(
            filename=files['filename'], url=files['url'])

    os.remove('temp.docx')
    os.remove('temp.pdf')

    return files['filename']


def parse_data(d: datetime.datetime, no: list) -> None:
    order = f"#AV{no[0]}{no[1]}{no[2]}"
    cells_text = [
        f"{d:%d}/{d:%m}/{d:%y} {d:%I}:{d:%M}:{d:%S} {d:%p}",
        f"{order}"
    ]
    if windows is True:
        document = Document('receipt_windows.docx')
    else:
        document = Document('receipt_linux.docx')
    table = document.tables[1].rows[0]

    for index, cell in enumerate(table.cells):
        paragraphs = cell.paragraphs
        for paragraph in paragraphs:
            run = paragraph.add_run()
            font = run.font
            if windows is True:
                font.name = WINDOWS[0]
                font.size = Pt(WINDOWS[1] + 1)
            else:
                font.name = LINUX[0]
                font.size = Pt(LINUX[1] + 1)
            font.bold = True
            run.text = cells_text[index]
    return document.save('temp.docx')


def parse_order(list_belanja: list, total_bayar: list, data: list) -> None:
    TEXT = ["SUBTOTAL", "PPN 10%", "TOTAL", "UANG", "CHANGE"]

    document = Document('temp.docx')
    table = document.tables[2]

    total = len(list_belanja) + 5
    for row in range(total):
        table.add_row()
    rows = table.rows

    # Tulis belanjaan pada struk belanja
    for index, belanja in enumerate(list_belanja):
        for i, cell in enumerate(rows[index].cells):
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                run = paragraph.add_run()
                font = run.font
                if windows is True:
                    font.name = WINDOWS[0]
                    font.size = Pt(WINDOWS[1])
                else:
                    font.name = LINUX[0]
                    font.size = Pt(LINUX[1])
                font.bold = True
                if i == 0:
                    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    run.text = list_belanja[index]
                elif i == 1:
                    run.text = rupiah_format(total_bayar[index])

    # Tulis SUBTOTAL, PPN, TOTAL, UANG, CHANGE/KEMBALIAN
    i = len(list_belanja) + 1
    j = i
    while i <= total:
        for index, cell in enumerate(rows[i].cells):
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                run = paragraph.add_run()
                font = run.font
                if windows is True:
                    font.name = WINDOWS[0]
                    font.size = Pt(WINDOWS[1])
                else:
                    font.name = LINUX[0]
                    font.size = Pt(LINUX[1])
                font.bold = True
                if index == 0:
                    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    run.text = TEXT[i - j]
                elif index == 1:
                    run.text = rupiah_format(data[i - j])
        i += 1

    return document.save('temp.docx')
